#!/usr/bin/env python3
"""
Convert Document to Word (.docx) with proper formatting
- All headings: Size 14, Bold, Black color
- Paragraphs: First letter indented, justified alignment
- One main topic per page
- DFD diagrams properly formatted
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
import re

def read_markdown(filepath):
    """Read markdown file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def format_heading(paragraph, text, level):
    """Format heading - Size 14, Bold, Black color, Times New Roman"""
    paragraph.clear()
    run = paragraph.add_run(text.lstrip('#').strip())
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    run.font.name = 'Times New Roman'
    
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    return paragraph

def create_styled_heading(doc, text, level):
    """Add a styled heading based on markdown level"""
    para = doc.add_paragraph()
    return format_heading(para, text, level)

def add_indented_paragraph(doc, text):
    """Add paragraph with first letter indentation and justified alignment"""
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Inches(0.5)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    
    # Format text color to black, Times New Roman, Size 12
    for run in para.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    return para

def add_table_of_contents(doc):
    """Add title page and table of contents with document info"""
    # Title Page
    title = doc.add_paragraph()
    title_run = title.add_run('DIET & NUTRITION TRACKING\nMOBILE APPLICATION')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.name = 'Times New Roman'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('TECHNICAL DOCUMENTATION')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
    subtitle_run.font.name = 'Times New Roman'
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(48)
    
    doc.add_page_break()
    
    # Table of Contents Page
    toc_title = doc.add_paragraph()
    toc_title_run = toc_title.add_run('TABLE OF CONTENTS')
    toc_title_run.font.size = Pt(14)
    toc_title_run.font.bold = True
    toc_title_run.font.color.rgb = RGBColor(0, 0, 0)
    toc_title_run.font.name = 'Times New Roman'
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(12)
    
    # Table of Contents entries
    toc_entries = [
        "Acknowledgement",
        "Abstract",
        "",
        "I. INTRODUCTION",
        "    1.1 Background",
        "    1.2 Scope of the Project",
        "    1.3 Objectives",
        "",
        "II. PROBLEM DEFINITION AND METHODOLOGY",
        "    2.1 Problem Definition",
        "    2.2 Existing System",
        "    2.3 Proposed System",
        "",
        "III. DEVELOPMENT PROCESS AND DESCRIPTION",
        "    3.1 Requirement Analysis and Specification",
        "        3.1.1 Frontend",
        "        3.1.2 Backend",
        "        3.1.3 Database",
        "        3.1.4 Hardware Requirements",
        "        3.1.5 Software Profile",
        "    3.2 System Design",
        "        3.2.1 Input Design",
        "        3.2.2 Output Design",
        "        3.2.3 Data Flow Diagram",
        "        3.2.4 Data Dictionary",
        "        3.2.5 Table Structure",
        "    3.3 Implementation",
        "        3.3.1 Front-End Implementation",
        "        3.3.2 Back-End Implementation",
        "        3.3.3 Database Implementation",
        "        3.3.4 Barcode Scanning Integration",
        "    3.4 Project Description",
        "        3.4.1 Dashboard Module",
        "        3.4.2 Food Search Module",
        "        3.4.3 Analytics Module",
        "        3.4.4 Settings Module",
        "    3.5 Testing",
        "        3.5.1 Unit Testing",
        "        3.5.2 Integration Testing",
        "        3.5.3 Functional Testing",
        "        3.5.4 User Interface (UI) Testing",
        "    3.6 Maintenance",
        "",
        "IV. CONCLUSION",
        "",
        "V. FUTURE ENHANCEMENT",
        "",
        "REFERENCES",
        "",
        "APPENDIX",
        "Source Code"
    ]
    
    for entry in toc_entries:
        if entry == "":
            doc.add_paragraph()
        else:
            para = doc.add_paragraph(entry)
            para.paragraph_format.left_indent = Inches(0.25 if entry.startswith("    ") else 0)
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_page_break()

def should_add_page_break(text):
    """Determine if page break should be added after heading"""
    main_sections = [
        'I. INTRODUCTION',
        'II. PROBLEM DEFINITION',
        'III. DEVELOPMENT PROCESS',
        'IV. CONCLUSION',
        'V. FUTURE ENHANCEMENT',
        'REFERENCES',
        'APPENDIX'
    ]
    
    for section in main_sections:
        if section.upper() in text.upper():
            return True
    return False

def process_markdown_content(filepath, output_path):
    """Convert content to Word document with proper formatting"""
    
    # Read file
    content = read_markdown(filepath)
    
    # Create Document
    doc = Document()
    
    # Set document to A4 size
    from docx.shared import Inches
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.7)  # A4 height
        section.page_width = Inches(8.27)   # A4 width
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title page and TOC
    add_table_of_contents(doc)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    last_heading_level = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            
            # Add page break for main sections
            if level == 1 and i > 0:
                doc.add_page_break()
            
            create_styled_heading(doc, heading_text, level)
            last_heading_level = level
            i += 1
            continue
        
        # Handle code blocks (for DFD diagrams and code)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines).strip()
                para = doc.add_paragraph(code_text)
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)
            
            i += 1
            continue
        
        # Handle tables (pipe-delimited)
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = [line]
            i += 1
            
            # Skip separator line
            if '---' in lines[i]:
                i += 1
            
            # Collect table rows
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            
            # Create table
            if len(table_lines) > 1:
                headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    hdr_cells[idx].text = header
                    # Format header
                    for paragraph in hdr_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Times New Roman'
                
                # Add body rows
                for row_line in table_lines[1:]:
                    cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                    if cells:
                        row_cells = table.add_row().cells
                        for idx, cell_text in enumerate(cells):
                            if idx < len(row_cells):
                                row_cells[idx].text = cell_text
                                for paragraph in row_cells[idx].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(11)
                                        run.font.name = 'Times New Roman'
            
            continue
        
        # Handle list items
        if line.strip().startswith('-'):
            list_text = line.strip()[1:].strip()
            para = doc.add_paragraph(list_text, style='List Bullet')
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip():
            text = line.strip()
            
            # Remove markdown formatting
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'__(.+?)__', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'_(.+?)_', r'\1', text)
            text = re.sub(r'\[(.+?)\]\((.+?)\)', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            
            add_indented_paragraph(doc, text)
        
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"✓ Document saved successfully: {output_path}")
    print(f"✓ Font: Times New Roman throughout")
    print(f"✓ Body text size: 12pt")
    print(f"✓ Headings: 14pt, Bold, Black")
    print(f"✓ Page layout: A4")
    print(f"✓ Table of Contents: Included")
    print(f"✓ Paragraphs: Justified with first-line indent")
    print(f"✓ Main sections on separate pages")

if __name__ == '__main__':
    # Check which markdown file exists
    markdown_files = [
        'Diet_Nutrition_Project_Report_Updated.md',
        'Diet_Nutrition_Project_Report.md',
        'Diet_Nutrition_Report_Final.md'
    ]
    
    markdown_file = None
    for mf in markdown_files:
        try:
            with open(mf, 'r', encoding='utf-8') as f:
                markdown_file = mf
                break
        except FileNotFoundError:
            continue
    
    if not markdown_file:
        print("❌ No markdown file found!")
        print(f"   Looked for: {', '.join(markdown_files)}")
    else:
        output_file = 'DIET_NUTRITION_TRACKING_APPLICATION.docx'
        
        print(f"📄 Converting {markdown_file}...")
        print(f"⚙️  Applying formatting:")
        print(f"   • Font: Times New Roman")
        print(f"   • Body text size: 12pt")
        print(f"   • Headings: 14pt, Bold, Black")
        print(f"   • Page layout: A4")
        print(f"   • Table of Contents: Included with complete structure")
        print(f"   • Data Flow Diagrams: Properly formatted")
        print(f"   • Paragraphs: Justified with first-line indent")
        print(f"   • Page breaks: After main sections")
        print()
        
        try:
            process_markdown_content(markdown_file, output_file)
            print()
            print(f"✅ Conversion complete!")
            print(f"📁 Output file: {output_file}")
        except Exception as e:
            print(f"❌ Error during conversion: {e}")
            import traceback
            traceback.print_exc()
